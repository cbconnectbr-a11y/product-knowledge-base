#!/usr/bin/env python3
"""
说明书内容提取模块

支持格式:
- PDF (主要): 使用 pypdf
- Word (次要): 使用 python-docx
- 其他格式暂不支持
"""

import logging
from io import BytesIO
from pathlib import Path
import lark_oapi as lark
from lark_oapi.api.drive.v1 import DownloadMediaRequest, DownloadFileRequest
import re

logger = logging.getLogger(__name__)


def extract_pdf_content(pdf_bytes: bytes) -> str:
    """
    从 PDF 提取文本内容

    Args:
        pdf_bytes: PDF 文件二进制内容

    Returns:
        提取的文本内容
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(pdf_bytes))

        # 提取所有页面文本
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())

        full_text = '\n\n'.join(text_parts)

        logger.info(f"PDF extraction: {len(reader.pages)} pages, {len(full_text)} chars")
        return full_text

    except Exception as e:
        logger.error(f"Failed to extract PDF content: {e}")
        return ''


def extract_word_content(docx_bytes: bytes) -> str:
    """
    从 Word 文档提取文本内容

    Args:
        docx_bytes: Word 文档二进制内容

    Returns:
        提取的文本内容
    """
    try:
        from docx import Document

        doc = Document(BytesIO(docx_bytes))

        # 提取所有段落
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        full_text = '\n'.join(paragraphs)

        logger.info(f"Word extraction: {len(paragraphs)} paragraphs, {len(full_text)} chars")
        return full_text

    except Exception as e:
        logger.error(f"Failed to extract Word content: {e}")
        return ''


def download_feishu_file(file_url: str, lark_client: lark.Client) -> bytes:
    """
    从飞书下载文件

    Args:
        file_url: 飞书文件链接 (格式: https://domain/file/{file_token})
        lark_client: 飞书客户端

    Returns:
        文件二进制内容

    Raises:
        Exception: 下载失败
    """
    # 从 URL 提取 file_token
    # 支持格式:
    # - https://domain/file/{file_token}
    # - https://domain/drive/folder/{folder_token}
    match = re.search(r'/file/([a-zA-Z0-9]+)', file_url)
    if not match:
        raise ValueError(f"Cannot extract file_token from URL: {file_url}")

    file_token = match.group(1)
    logger.info(f"Extracted file_token: {file_token}")

    def _read(response) -> bytes:
        if not response.success():
            return b''
        f = getattr(response, "file", None)
        if f is not None:
            try:
                data = f.read()
                if data:
                    return data
            except Exception as e:
                logger.error(f"读取 response.file 失败: {e}")
        raw = getattr(response, "raw", None)
        if raw is not None and getattr(raw, "content", None):
            return raw.content
        return b''

    # 云空间 /file/ 文件用 file.download；文档素材用 media.download。优先 file，回退 media。
    content = b''
    try:
        content = _read(lark_client.drive.v1.file.download(
            DownloadFileRequest.builder().file_token(file_token).build()))
        if content:
            logger.info(f"file.download 成功: {len(content)} bytes")
    except Exception as e:
        logger.warning(f"file.download 异常，尝试 media.download: {e}")

    if not content:
        try:
            content = _read(lark_client.drive.v1.media.download(
                DownloadMediaRequest.builder().file_token(file_token).build()))
            if content:
                logger.info(f"media.download 成功: {len(content)} bytes")
        except Exception as e:
            logger.warning(f"media.download 异常: {e}")

    if not content:
        raise Exception("下载 0 字节 —— 文件为空或应用无该文件访问权")
    return content


def extract_manual_content(manual_files: dict, lark_client: lark.Client) -> str:
    """
    提取说明书内容 (统一入口)

    Args:
        manual_files: 说明书字段 (格式: {'link': '...', 'text': 'filename'})
        lark_client: 飞书客户端

    Returns:
        提取的文本内容
    """
    if not manual_files or not isinstance(manual_files, dict):
        return ''

    file_url = manual_files.get('link')
    filename = manual_files.get('text', '')

    if not file_url:
        return ''

    try:
        # 下载文件
        file_bytes = download_feishu_file(file_url, lark_client)

        # 根据文件扩展名选择提取方法
        ext = Path(filename).suffix.lower() if filename else ''

        if ext == '.pdf':
            content = extract_pdf_content(file_bytes)
        elif ext in ['.docx', '.doc']:
            content = extract_word_content(file_bytes)
        else:
            logger.warning(f"Unsupported file format: {ext} ({filename})")
            return ''

        return content

    except Exception as e:
        logger.error(f"Failed to extract manual content from {filename}: {e}")
        return ''


if __name__ == '__main__':
    # 测试代码
    import os
    from dotenv import load_dotenv

    load_dotenv()

    logging.basicConfig(level=logging.INFO)

    # 创建飞书客户端
    client = lark.Client.builder() \
        .app_id(os.environ.get('FEISHU_APP_ID')) \
        .app_secret(os.environ.get('FEISHU_APP_SECRET')) \
        .build()

    # 测试提取
    test_manual = {
        'link': 'https://cgokyyxlsh.feishu.cn/file/BP0jbkQHZot0D7xiUsVcRqojn6e',
        'text': '水龙头通用说明书葡语.pdf'
    }

    content = extract_manual_content(test_manual, client)

    if content:
        print(f'\n✅ 提取成功: {len(content)} 字符')
        print(f'\n内容预览:')
        print('=' * 70)
        print(content[:500])
    else:
        print('\n❌ 提取失败')
